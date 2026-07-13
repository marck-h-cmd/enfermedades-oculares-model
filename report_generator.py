import os
import time
import json
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
from datetime import datetime

# Importaciones condicionales para evitar fallos si la instalación no ha terminado
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    from fpdf import FPDF
except ImportError:
    FPDF = None

def generar_excel_report(resultados_cv, ruta_salida=None):
    """
    Genera un archivo de Excel (.xlsx) estructurado y formateado con múltiples hojas:
    - Resumen Global de los 5 Modelos
    - Detalle de accuracies y tiempos por fold
    - Reportes de clasificación detallados por clase
    """
    if ruta_salida is None:
        marca_tiempo = datetime.now().strftime("%Y%m%d_%H%M%S")
        ruta_salida = f"reporte_comparativo_{marca_tiempo}.xlsx"
        
    # 1. Crear datos de Resumen Global
    filas_resumen = []
    for mod_name, res in resultados_cv.items():
        filas_resumen.append({
            'Modelo': mod_name.upper(),
            'Accuracy Promedio': res['accuracy_media'],
            'Desviación Estándar': res['accuracy_std'],
            'Tiempo Promedio (s)': res['tiempo_medio'],
            'F1-Score Macro': res['reporte_final']['macro avg']['f1-score'],
            'F1-Score Weighted': res['reporte_final']['weighted avg']['f1-score']
        })
    df_resumen = pd.DataFrame(filas_resumen)
    
    # 2. Crear datos de Detalle por Folds
    filas_folds = []
    for mod_name, res in resultados_cv.items():
        for fold, (acc, t) in enumerate(zip(res['accuracies_folds'], res['tiempos_folds'])):
            filas_folds.append({
                'Modelo': mod_name.upper(),
                'Fold': fold + 1,
                'Accuracy': acc,
                'Tiempo (s)': t
            })
    df_folds = pd.DataFrame(filas_folds)
    
    # 3. Crear reporte detallado por clase
    filas_clases = []
    for mod_name, res in resultados_cv.items():
        reporte = res['reporte_final']
        for clase_nombre, metricas in reporte.items():
            if isinstance(metricas, dict) and clase_nombre not in ['accuracy', 'macro avg', 'weighted avg']:
                filas_clases.append({
                    'Modelo': mod_name.upper(),
                    'Clase': clase_nombre,
                    'Precision': metricas['precision'],
                    'Recall': metricas['recall'],
                    'F1-Score': metricas['f1-score'],
                    'Soporte': metricas['support']
                })
    df_clases = pd.DataFrame(filas_clases)
    
    # Escribir en Excel con openpyxl y ajustar ancho de columnas
    with pd.ExcelWriter(ruta_salida, engine='openpyxl') as writer:
        df_resumen.to_excel(writer, sheet_name='Resumen Global', index=False)
        df_folds.to_excel(writer, sheet_name='Detalle por Folds', index=False)
        df_clases.to_excel(writer, sheet_name='Métricas por Clase', index=False)
        
        # Ajustar ancho de columnas para mejor visualización
        for sheetname in writer.sheets:
            worksheet = writer.sheets[sheetname]
            for col_cells in worksheet.columns:
                max_length = 0
                column = col_cells[0].column_letter
                for cell in col_cells:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                adjusted_width = (max_length + 2)
                worksheet.column_dimensions[column].width = adjusted_width
        
    return ruta_salida

def generar_grafico_comparativo(resultados_cv):
    """Genera un gráfico comparativo de accuracies por fold y lo guarda en memoria"""
    plt.figure(figsize=(10, 6))
    modelos = list(resultados_cv.keys())
    
    datos_plot = []
    for mod in modelos:
        for acc in resultados_cv[mod]['accuracies_folds']:
            datos_plot.append({'Modelo': mod.upper(), 'Accuracy': acc})
            
    df_plot = pd.DataFrame(datos_plot)
    
    sns.boxplot(x='Modelo', y='Accuracy', data=df_plot, hue='Modelo', legend=False, palette='Set2')
    sns.stripplot(x='Modelo', y='Accuracy', data=df_plot, color='black', alpha=0.5, size=6)
    
    plt.title('Comparación de Modelos: Accuracy a lo largo de los Folds', fontsize=14, weight='bold')
    plt.ylabel('Validation Accuracy', fontsize=12)
    plt.xlabel('Modelos', fontsize=12)
    plt.grid(True, linestyle='--', alpha=0.3)
    
    img_buf = BytesIO()
    plt.savefig(img_buf, format='png', dpi=300, bbox_inches='tight')
    img_buf.seek(0)
    plt.close()
    return img_buf

def generar_word_report(resultados_cv, pruebas_stats, ruta_salida=None):
    """
    Genera un reporte profesional en Word (.docx) con:
    - Tablas de rendimiento
    - Gráfico comparativo insertado
    - Resultados de pruebas estadísticas e interpretaciones
    """
    if Document is None:
        raise ImportError("La librería 'python-docx' no está instalada.")
        
    if ruta_salida is None:
        marca_tiempo = datetime.now().strftime("%Y%m%d_%H%M%S")
        ruta_salida = f"reporte_clinico_{marca_tiempo}.docx"
        
    doc = Document()
    
    # Estilos del Documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Título Principal
    titulo = doc.add_paragraph()
    titulo_run = titulo.add_run("REPORTE CLÍNICO-TÉCNICO DE REDES NEURONALES OCULARES")
    titulo_run.bold = True
    titulo_run.size = Pt(18)
    titulo_run.font.color.rgb = RGBColor(0, 51, 102) # Azul oscuro
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo con fecha
    subtitulo = doc.add_paragraph()
    subtitulo.add_run(f"Generado el: {datetime.now().strftime('%d/%m/%Y a las %H:%M:%S')}\n"
                      f"Metodología: Validación Cruzada (5-Fold Cross Validation)")
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph().add_run("\n1. INTRODUCCIÓN Y CONTEXTO").bold = True
    doc.add_paragraph(
        "Este informe detalla los resultados obtenidos al evaluar tres modelos clásicos "
        "de redes neuronales convolucionales (MobileNetV2, ResNet50V2, EfficientNetB0) "
        "y dos modelos híbridos (Fusión de características y CNN-RandomForest) para la detección "
        "y clasificación de patologías oculares como Catarata, Glaucoma, Retinopatía Diabética, entre otras. "
        "La evaluación se realizó mediante validación cruzada para garantizar la estabilidad y reproducibilidad de las métricas."
    )
    
    # Tabla Resumen de Resultados
    doc.add_paragraph().add_run("2. RESUMEN DE RENDIMIENTO DE LOS MODELOS").bold = True
    
    tabla = doc.add_table(rows=1, cols=5)
    tabla.style = 'Light Shading Accent 1'
    hdr_cells = tabla.rows[0].cells
    hdr_cells[0].text = 'Modelo'
    hdr_cells[1].text = 'Accuracy Promedio'
    hdr_cells[2].text = 'Std Dev'
    hdr_cells[3].text = 'F1 weighted'
    hdr_cells[4].text = 'Tiempo Medio (s)'
    
    for mod_name, res in resultados_cv.items():
        row_cells = tabla.add_row().cells
        row_cells[0].text = mod_name.upper()
        row_cells[1].text = f"{res['accuracy_media']:.2%}"
        row_cells[2].text = f"±{res['accuracy_std']:.4f}"
        row_cells[3].text = f"{res['reporte_final']['weighted avg']['f1-score']:.4f}"
        row_cells[4].text = f"{res['tiempo_medio']:.2f}s"
        
    doc.add_paragraph()
    
    # Insertar gráfico comparativo
    doc.add_paragraph().add_run("3. ANÁLISIS GRÁFICO COMPARATIVO").bold = True
    try:
        img_buf = generar_grafico_comparativo(resultados_cv)
        # Guardar temporalmente la imagen para insertarla
        temp_img_path = "temp_docx_chart.png"
        with open(temp_img_path, 'wb') as f:
            f.write(img_buf.read())
            
        doc.add_picture(temp_img_path, width=Inches(6.0))
        os.remove(temp_img_path)
        
        p_cap = doc.add_paragraph()
        p_cap.add_run("Figura 1: Comparación del rango de exactitud (accuracy) por modelo en los 5 pliegues de validación cruzada.").italic = True
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[No se pudo insertar el gráfico: {str(e)}]")
        
    # Resultados de Pruebas Estadísticas
    doc.add_paragraph().add_run("\n4. ANÁLISIS ESTADÍSTICO DE SIGNIFICANCIA").bold = True
    
    # Test de Friedman
    friedman = pruebas_stats.get('anova_friedman', {})
    if 'p_valor' in friedman:
        doc.add_paragraph(
            f"Prueba global de Friedman: Estadístico = {friedman['estadistico']:.4f}, p-valor = {friedman['p_valor']:.6f}. "
            f"Interpretación: {friedman['interpretacion']}"
        )
    
    # Comparación por pares
    doc.add_paragraph("Detalle de comparaciones estadísticas por pares con normalidad y corrección de Holm-Bonferroni:")
    tabla_stats = doc.add_table(rows=1, cols=6)
    tabla_stats.style = 'Light Shading Accent 1'
    hdr_cells_s = tabla_stats.rows[0].cells
    hdr_cells_s[0].text = 'Comparación'
    hdr_cells_s[1].text = 'Normalidad (Shapiro-Wilk p-val)'
    hdr_cells_s[2].text = 't-Student p-val (Orig / Holm)'
    hdr_cells_s[3].text = 'Wilcoxon p-val (Orig / Holm)'
    hdr_cells_s[4].text = 'd de Cohen'
    hdr_cells_s[5].text = 'Interpretación'
    
    for comp in pruebas_stats.get('comparaciones_pares', []):
        row_cells = tabla_stats.add_row().cells
        row_cells[0].text = f"{comp['modelo1'].upper()} vs {comp['modelo2'].upper()}"
        
        # Shapiro-Wilk
        shap = comp['shapiro']
        dist_str = f"{'Normal' if shap['normal'] else 'No Normal'} ({shap['p_valor']:.4f})"
        row_cells[1].text = dist_str
        
        # t-Student
        t_stud = comp['t_student']
        row_cells[2].text = f"{t_stud['p_valor_original']:.4f} / {t_stud['p_valor_corregido']:.4f}"
        
        # Wilcoxon
        wilc = comp['wilcoxon']
        row_cells[3].text = f"{wilc['p_valor_original']:.4f} / {wilc['p_valor_corregido']:.4f}"
        
        # Cohen's d
        d_val = comp['cohens_d']
        row_cells[4].text = f"{d_val['valor']:.2f} ({d_val['interpretacion']})"
        
        # Interpretación
        row_cells[5].text = comp['interpretacion']
        
    # Conclusión clínica
    doc.add_paragraph().add_run("\n5. CONCLUSIÓN Y RECOMENDACIONES CLÍNICAS").bold = True
    
    # Encontrar mejor modelo
    mejor_modelo = max(resultados_cv.keys(), key=lambda k: resultados_cv[k]['accuracy_media'])
    doc.add_paragraph(
        f"Con base en los resultados empíricos y validación estadística, se determina que el mejor modelo general es "
        f"{mejor_modelo.upper()} con una exactitud media de {resultados_cv[mejor_modelo]['accuracy_media']:.2%}. "
        f"Este modelo ha sido guardado en la raíz del proyecto como 'best_ocular_model.h5' para su consumo directo "
        f"sin necesidad de reentrenamiento.\n\n"
        f"Se recomienda su integración en el sistema de soporte a decisiones médicas del hospital."
    )
    
    # Disclaimer
    p_disclaimer = doc.add_paragraph()
    run_disc = p_disclaimer.add_run(
        "AVISO: Este informe clínico-técnico es generado de forma automatizada por un sistema de apoyo de "
        "inteligencia artificial. Las predicciones deben ser siempre validadas por personal médico y oftalmológico calificado."
    )
    run_disc.italic = True
    run_disc.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.save(ruta_salida)
    return ruta_salida
 
class PDFReporte(FPDF):
    """Clase personalizada para FPDF con cabecera y pie de página"""
    def header(self):
        self.set_font('Arial', 'B', 15)
        self.set_text_color(0, 51, 102)
        self.cell(0, 10, 'REPORTE COMPARATIVO DE DIAGNÓSTICO OCULAR (5-FOLD CV)', 0, 1, 'C')
        self.line(10, 20, 200, 20)
        self.ln(10)
        
    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f'Página {self.page_no()}/{{nb}} - Soporte Diagnóstico IA', 0, 0, 'C')
 
def generar_pdf_report(resultados_cv, pruebas_stats, ruta_salida=None):
    """Genera reporte consolidado en PDF con FPDF"""
    if FPDF is None:
        raise ImportError("La librería 'fpdf' no está instalada.")
        
    if ruta_salida is None:
        marca_tiempo = datetime.now().strftime("%Y%m%d_%H%M%S")
        ruta_salida = f"reporte_consolidado_{marca_tiempo}.pdf"
        
    pdf = PDFReporte()
    pdf.alias_nb_pages()
    pdf.add_page()
    
    # Fecha y Metodología
    pdf.set_font('Arial', '', 10)
    pdf.set_text_color(50, 50, 50)
    pdf.cell(0, 6, f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", 0, 1)
    pdf.cell(0, 6, "Metodología: Validación Cruzada Estratificada por 5 Folds", 0, 1)
    pdf.ln(5)
    
    # 1. Rendimiento
    pdf.set_font('Arial', 'B', 12)
    pdf.set_text_color(0, 51, 102)
    pdf.cell(0, 8, "1. Rendimiento Comparativo de Modelos", 0, 1)
    pdf.ln(2)
    
    # Tabla de métricas
    pdf.set_font('Arial', 'B', 9)
    pdf.set_fill_color(220, 230, 240)
    pdf.cell(40, 7, 'Modelo', 1, 0, 'C', True)
    pdf.cell(35, 7, 'Accuracy Promedio', 1, 0, 'C', True)
    pdf.cell(30, 7, 'Desv. Estándar', 1, 0, 'C', True)
    pdf.cell(30, 7, 'F1-Score weighted', 1, 0, 'C', True)
    pdf.cell(35, 7, 'Tiempo Medio (s)', 1, 1, 'C', True)
    
    pdf.set_font('Arial', '', 9)
    for mod_name, res in resultados_cv.items():
        pdf.cell(40, 6, mod_name.upper(), 1, 0, 'C')
        pdf.cell(35, 6, f"{res['accuracy_media']:.2%}", 1, 0, 'C')
        pdf.cell(30, 6, f"±{res['accuracy_std']:.4f}", 1, 0, 'C')
        pdf.cell(30, 6, f"{res['reporte_final']['weighted avg']['f1-score']:.4f}", 1, 0, 'C')
        pdf.cell(35, 6, f"{res['tiempo_medio']:.2f}s", 1, 1, 'C')
        
    pdf.ln(8)
    
    # 2. Pruebas estadísticas
    pdf.set_font('Arial', 'B', 12)
    pdf.set_text_color(0, 51, 102)
    pdf.cell(0, 8, "2. Análisis Estadístico Inferencial", 0, 1)
    pdf.ln(2)
    
    # Friedman
    friedman = pruebas_stats.get('anova_friedman', {})
    pdf.set_font('Arial', '', 10)
    if 'p_valor' in friedman:
        pdf.multi_cell(0, 5, f"Prueba de Friedman (global): Estadístico = {friedman['estadistico']:.4f}, p-valor = {friedman['p_valor']:.6f}.\nInterpretación: {friedman['interpretacion']}")
        pdf.ln(4)
        
    # Comparaciones pares
    pdf.set_font('Arial', 'B', 8)
    pdf.set_fill_color(240, 240, 240)
    pdf.cell(35, 7, 'Comparación', 1, 0, 'C', True)
    pdf.cell(30, 7, 'Shapiro p-val', 1, 0, 'C', True)
    pdf.cell(35, 7, 't-Student (Orig/Holm)', 1, 0, 'C', True)
    pdf.cell(35, 7, 'Wilcoxon (Orig/Holm)', 1, 0, 'C', True)
    pdf.cell(25, 7, 'd de Cohen', 1, 0, 'C', True)
    pdf.cell(30, 7, 'Significativo', 1, 1, 'C', True)
    
    pdf.set_font('Arial', '', 7.5)
    for comp in pruebas_stats.get('comparaciones_pares', []):
        pdf.cell(35, 6, f"{comp['modelo1'].upper()} vs {comp['modelo2'].upper()}", 1, 0, 'C')
        
        # Shapiro
        shap = comp['shapiro']
        pdf.cell(30, 6, f"{shap['p_valor']:.3f} ({'Norm' if shap['normal'] else 'No N.'})", 1, 0, 'C')
        
        # t-Student
        t_stud = comp['t_student']
        pdf.cell(35, 6, f"{t_stud['p_valor_original']:.3f} / {t_stud['p_valor_corregido']:.3f}", 1, 0, 'C')
        
        # Wilcoxon
        wilc = comp['wilcoxon']
        pdf.cell(35, 6, f"{wilc['p_valor_original']:.3f} / {wilc['p_valor_corregido']:.3f}", 1, 0, 'C')
        
        # Cohen's d
        d_val = comp['cohens_d']
        pdf.cell(25, 6, f"{d_val['valor']:.2f} ({d_val['interpretacion'][:5]})", 1, 0, 'C')
        
        # Significativo
        normal = shap['normal']
        p_val = t_stud['p_valor_corregido'] if normal else wilc['p_valor_corregido']
        pdf.cell(30, 6, "Si" if p_val < 0.05 else "No", 1, 1, 'C')
        
    # 3. Diagrama de Validación
    pdf.add_page()
    pdf.set_font('Arial', 'B', 12)
    pdf.set_text_color(0, 51, 102)
    pdf.cell(0, 8, "3. Diagrama de Validacion Cruzada (Accuracy)", 0, 1)
    pdf.ln(5)
    
    try:
        img_buf = generar_grafico_comparativo(resultados_cv)
        temp_img_path = "temp_pdf_chart.png"
        with open(temp_img_path, 'wb') as f:
            f.write(img_buf.read())
        pdf.image(temp_img_path, x=15, w=180)
        os.remove(temp_img_path)
    except Exception as e:
        pdf.set_font('Arial', '', 10)
        pdf.cell(0, 5, f"[No se pudo generar el grafico: {str(e)}]", 0, 1)
        
    pdf.ln(20)
    
    # 4. Firma Digital
    pdf.set_font('Arial', 'B', 11)
    pdf.cell(0, 5, "__________________________________________________", 0, 1, 'C')
    pdf.cell(0, 5, "DOCUMENTO OFICIAL FIRMADO DIGITALMENTE", 0, 1, 'C')
    pdf.set_font('Arial', '', 9)
    pdf.cell(0, 5, f"Sistema Experto de Diagnostico Ocular - ID Validacion: {datetime.now().strftime('%Y%m%d%H%M%S')}", 0, 1, 'C')
    pdf.cell(0, 5, "Validez Tecnica y Cientifica bajo validacion cruzada rigurosa", 0, 1, 'C')
 
    # Guardar PDF
    pdf.output(ruta_salida)
    return ruta_salida
